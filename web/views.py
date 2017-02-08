# -*- coding: UTF-8 -*-
from django.shortcuts import render_to_response, redirect
from django.http import HttpResponse
from django.template import RequestContext
from models import Diary, Month, Money
from forms import DiaryForm, MoneyForm, LoginForm
from django.core.exceptions import ObjectDoesNotExist
from django.utils import timezone
from django.utils.timezone import localtime
import StringIO
from docx import *
from docx.shared import Inches
import xlsxwriter
from django.contrib.auth import authenticate, login
from django.contrib.auth.decorators import login_required

# 瀏覽日誌
@login_required
def diary(request, month):
        time_year = int(month)/100
        time_month = int(month)%100
        diaries = Diary.objects.filter(time__year=time_year, time__month=time_month).order_by("-id")
        return render_to_response('diary.html', {'diaries': diaries, 'month':month}, context_instance=RequestContext(request))


def diary_add(request):
        if request.method == 'POST':
                form = DiaryForm(request.POST)
                if form.is_valid():
                        form.save()
                        year = localtime(timezone.now()).year
                        month =  localtime(timezone.now()).month
                        try:
                                themonth = Month.objects.get(date=year*100+month)
                        except ObjectDoesNotExist:
                                themonth = Month(date=year*100+month)
                                themonth.save()
                        return redirect("/diary")
        else:
                form = DiaryForm()
        return render_to_response('form.html',{'form': form}, context_instance=RequestContext(request))    

@login_required      
def home(request):
        months = Month.objects.all().order_by("-id")
        return render_to_response('home.html', {'months': months}, context_instance=RequestContext(request))      

@login_required
def diary_word(request, month):
        document = Document()
        docx_title="Diary-"+str(timezone.localtime(timezone.now()).date())+".docx"

        time_year = int(month)/100
        time_month = int(month)%100
        diaries = Diary.objects.filter(time__year=time_year, time__month=time_month).order_by("-id")
        paragraph = document.add_paragraph(u'我的日誌：'+month)
        table = document.add_table(rows=1, cols=2)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'日期'
        hdr_cells[1].text = u'內容'
        for diary in diaries:
                row_cells = table.add_row().cells
                row_cells[0].text = str(timezone.localtime(diary.time).strftime("%b %d %Y %H:%M:%S"))
                row_cells[1].text = diary.memo

        # Prepare document for download
        # -----------------------------
        f = StringIO.StringIO()
        document.save(f)
        length = f.tell()
        f.seek(0)
        response = HttpResponse(
                f.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=' + docx_title
        response['Content-Length'] = length

        return response   
      
@login_required      
def money(request, month):
        time_year = int(month)/100
        time_month = int(month)%100
        moneys = Money.objects.filter(time__year=time_year, time__month=time_month).order_by("-id")
        return render_to_response('money.html', {'moneys': moneys, 'month':month}, context_instance=RequestContext(request))

@login_required     
def money_add(request):
        if request.method == 'POST':
                form = MoneyForm(request.POST)
                if form.is_valid():
                        form.save()
                        year = localtime(timezone.now()).year
                        month =  localtime(timezone.now()).month
                        try:
                                themonth = Month.objects.get(date=year*100+month)
                        except ObjectDoesNotExist:
                                themonth = Month(date=year*100+month)
                                themonth.save()
                        return redirect("/money/"+str(themonth.date))
        else:
                form = MoneyForm()
        return render_to_response('form.html',{'form': form}, context_instance=RequestContext(request))
      
@login_required      
def money_excel(request, month):
        time_year = int(month)/100
        time_month = int(month)%100
        moneys = Money.objects.filter(time__year=time_year, time__month=time_month).order_by("-id")
        output = StringIO.StringIO()
        workbook = xlsxwriter.Workbook(output)
        worksheet = workbook.add_worksheet(month)
        worksheet.write(0,0, u"項目")
        worksheet.write(0,1, u"類別")
        worksheet.write(0,2, u"費用")
        worksheet.write(0,3, u"時間")
        counter = 1
        for money in moneys:
                worksheet.write(counter,0, money.item)
                worksheet.write(counter,1, money.kind)
                worksheet.write(counter,2, money.price)
                worksheet.write(counter,3, str(localtime(money.time).strftime("%b %d %Y %H:%M:%S")))
                counter = counter + 1
        workbook.close()
        # xlsx_data contains the Excel file
        response = HttpResponse(content_type='application/vnd.ms-excel')
        response['Content-Disposition'] = 'attachment; filename=Money-'+str(localtime(timezone.now()).date())+'.xlsx'
        xlsx_data = output.getvalue()
        response.write(xlsx_data)
        return response
      
# 使用者登入功能
def user_login(request):
        message = ""
        if request.method == "POST":
                        form = LoginForm(request.POST)
                        if form.is_valid():
                                username = request.POST['username']
                                password = request.POST['password']
                                user = authenticate(username=username, password=password)
                                if user is not None:
                                        if user.is_active:
                                                # 登入成功，導到大廳
                                                login(request, user)
                                                return redirect('/home')
                                        else:
                                                message = "無效的帳號或密碼!"
        else:
                form = LoginForm()
        return render_to_response('login.html', {'message': message, 'form': form}, context_instance=RequestContext(request))
